from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

from aker_core.exports import MemoContextBuilder, to_word


@pytest.fixture
def memo_context(image_payload: dict[str, str]) -> dict[str, object]:
    builder = MemoContextBuilder()
    memo = builder.build(
        msa={
            "msa_id": "BOI",
            "name": "Boise",
            "as_of": "2024-10-01",
            "pillar_scores": {
                "weighted_0_5": 4.2,
                "risk_multiplier": 0.95,
            },
        },
        market_tables={
            "supply": [
                {"metric_label": "Permit Velocity", "value_raw": 120, "unit": "bps", "normalized_0_100": 82},
            ],
            "jobs": [
                {"metric_label": "STEM Hiring", "value_raw": 18, "unit": "%", "normalized_0_100": 88},
            ],
            "urban": [
                {"metric_label": "Transit Score", "value_raw": 65, "unit": "idx", "normalized_0_100": 70},
            ],
            "outdoors": [
                {"metric_label": "Trail Access", "value_raw": 42, "unit": "idx", "normalized_0_100": 77},
            ],
        },
        risk={
            "multipliers": [
                {"peril": "wildfire", "multiplier": 1.1, "exit_cap_bps_delta": 15, "contingency_pct_delta": 0.05},
            ],
            "exit_cap_bps_delta": 15,
            "contingency_pct_delta": 0.05,
        },
        data_vintage={"BLS": "2024-09", "HUD": "2024-08"},
        run_id="run-123",
        git_sha="abc1234",
        created_at=datetime(2024, 10, 5, 12, 30, 0),
        images=image_payload,
        asset={
            "asset_id": "asset-001",
            "fit_score": 87,
            "flags": [
                {
                    "severity": "medium",
                    "code": "AMENITY",
                    "message": "Amenity deficit",
                    "observed": "Pool offline",
                    "target": "Pool refreshed",
                }
            ],
        },
        ops={
            "reputation_idx": 72,
            "nps_series": [
                {"period": "2024-Q2", "score": 35},
            ],
            "reviews_series": [
                "Tenant email john.doe@example.com left number 555-123-9876",
            ],
        },
        state_pack={
            "code": "ID",
            "changes": [
                {"section": "Water Rights", "setting": "allocation", "current": "120 gpd", "proposed": "140 gpd"},
            ],
        },
    )
    context = dict(memo.data)
    context["_memo_meta"] = memo.metadata
    return context


def test_to_word_renders_complete_memo(tmp_path: Path, memo_context: dict[str, object]) -> None:
    output_path = to_word(memo_context, outdir=tmp_path)

    assert output_path.exists()

    with ZipFile(output_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        assert "{{" not in document_xml
        assert "{%" not in document_xml
        assert "Executive Summary" in document_xml
        assert "Appendix A – Data sources" in document_xml

        media_entries = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert len(media_entries) >= 1

    document = Document(output_path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    normalised = [text.replace("2024-10-05T12:30:00", "<CREATED_AT>") for text in paragraphs]
    golden_path = Path(__file__).parent / "golden" / "ic_memo_paragraphs.txt"
    expected = golden_path.read_text(encoding="utf-8").splitlines()
    assert normalised[: len(expected)] == expected


def test_to_word_missing_images_warns(tmp_path: Path, memo_context: dict[str, object], caplog: pytest.LogCaptureFixture) -> None:
    memo_context = dict(memo_context)
    images = dict(memo_context["images"])
    images["pillar_bars"] = "non-existent.png"
    memo_context["images"] = images
    caplog.set_level("WARNING")

    output_path = to_word(memo_context, outdir=tmp_path)
    assert output_path.exists()

    doc = Document(output_path)
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "[IMAGE MISSING: pillar_bars]" in texts


def test_memo_context_builder_redacts_reviews(image_payload: dict[str, str]) -> None:
    builder = MemoContextBuilder()
    payload = builder.build(
        msa={
            "msa_id": "BOI",
            "name": "Boise",
            "as_of": "2024-10-01",
            "pillar_scores": {"weighted_0_5": 4.2, "risk_multiplier": 0.95},
        },
        market_tables={"supply": [], "jobs": [], "urban": [], "outdoors": []},
        risk={"multipliers": []},
        data_vintage={},
        run_id="run-123",
        git_sha="abc1234",
        created_at=datetime.now(timezone.utc),
        images=image_payload,
        ops={"reputation_idx": 70, "reviews_series": ["Email me at jane@corp.com or call 415-555-9988"]},
    )

    assert payload.metadata["redactions"] == 2
    assert "[REDACTED EMAIL]" in payload.data["ops"]["reviews_series"][0]
    assert "[REDACTED PHONE]" in payload.data["ops"]["reviews_series"][0]
